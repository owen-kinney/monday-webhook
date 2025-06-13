import requests
from datetime import datetime, timedelta, timezone



# === CONFIG ===
API_KEY = 'eyJhbGciOiJIUzI1NiJ9.eyJ0aWQiOjQzMDgzOTgxNCwiYWFpIjoxMSwidWlkIjo2Mjg4Mjg0MywiaWFkIjoiMjAyNC0xMC0zMVQxODoyNDo1NS4wMDBaIiwicGVyIjoibWU6d3JpdGUiLCJhY3RpZCI6MjAwMTk5MjMsInJnbiI6InVzZTEifQ.OwrhlolZ6uLUI2LQQ8pjA1DRPCc61kx_n3a3j6lM09E'  # Replace with your actual key

BOARD_ID = 8864427840

headers = {
    "Authorization": API_KEY
}

query = f"""
{{
  boards(ids: [{BOARD_ID}]) {{
    name
    items_page(limit: 100) {{
      items {{
        name
        group {{
          title
        }}
        column_values(ids: ["owner"]) {{
          text
        }}
        updates {{
          text_body
          created_at
        }}
      }}
    }}
  }}
}}
"""

response = requests.post(
    url="https://api.monday.com/v2",
    json={"query": query},
    headers=headers
)

data = response.json()

# ✅ Date filter: past 7 days
now = datetime.now(timezone.utc)
seven_days_ago = now - timedelta(days=7)

if 'data' in data and data['data']['boards']:
    board = data['data']['boards'][0]
    grouped = {}

    for item in board['items_page']['items']:
        group_title = item['group']['title']
        item_name = item['name']
        owners_text = item['column_values'][0]['text'] if item['column_values'] else 'None'

        # Filter updates from last 7 days
        recent_updates = []
        for update in item.get('updates', []):
            update_time = datetime.fromisoformat(update['created_at'].replace('Z', '+00:00'))
            update_text = update['text_body'].strip()
            if update_time >= seven_days_ago and "opp weekly status update" in update_text.lower():
                recent_updates.append(update['text_body'].strip())

        weekly_update = recent_updates[0] if recent_updates else "Nothing to report."
        grouped.setdefault(group_title, []).append((item_name, owners_text, weekly_update))

    # ✅ Output
   # print(f"\nBoard: {board['name']}\n")
    #for group, items in grouped.items():
     #   print(group)
      #  for item_name, owners, update in items:
       #     print(f"  - {item_name}")
        #    print(f"    • Owners: {owners}")
         #   print(f"    • {update}")
        #print()
#else:
 #   print("❌ Could not retrieve items.")

from html import escape

html =""

for group, items in grouped.items():
    html += f"<h2>{escape(group)}</h2><ul>"
    for item_name, owners, update in items:
        html += "<li>"
        html += f"<strong>{escape(item_name)}</strong><br>"
        html += f"• <em>Owners:</em> {escape(owners)}<br>"
        html += f"• <em></em> {escape(update)}"
        html += "</li>"
    html += "</ul><br>"

import msal

# === MSAL Setup ===
client_id = "de423635-3ca8-40e4-bd82-7413dfa46174"
authority = "https://login.microsoftonline.com/a4086c0d-45ae-4eaa-ab46-58f521bb203d"# Or use your tenant ID if needed
scopes = ["Mail.Send"]

# === Email Info ===
sender_email = "owen.kinney@cobaltsp.com"  # Your email
recipient_email = "tyler.hoffman@cobaltsp.com"  # Recipient
today_str = datetime.today().strftime("%m/%d/%y")
subject = f"Weekly OPP Status Update {today_str}"

# === MSAL Interactive Login ===
app = msal.PublicClientApplication(client_id=client_id, authority=authority)
accounts = app.get_accounts()

if accounts:
    result = app.acquire_token_silent(scopes, account=accounts[0])
else:
    result = app.acquire_token_interactive(scopes=scopes)

if "access_token" in result:
    headers = {
        "Authorization": f"Bearer {result['access_token']}",
        "Content-Type": "application/json"
    }

    email_data = {
        "message": {
            "subject": subject,
            "body": {
                "contentType": "HTML",
                "content": html  # Assuming `html` is your already-formatted summary
            },
            "toRecipients": [
                {"emailAddress": {"address": recipient_email}}
            ]
        },
        "saveToSentItems": "true"
    }

    response = requests.post(
        "https://graph.microsoft.com/v1.0/me/sendMail",
        headers=headers,
        json=email_data
    )

    if response.status_code == 202:
        print("✅ Email sent successfully!")
    else:
        print(f"❌ Failed to send email: {response.status_code} - {response.text}")
else:
    print("❌ Authentication failed:", result)

from docx import Document
from datetime import datetime
import os

# === Prep ===
today_str = datetime.today().strftime("%m/%d/%y")     # for subject
safe_date = datetime.today().strftime("%m-%d-%y")     # for filename
subject = f"Weekly OPP Status Update {today_str}"

# === Create document ===
doc = Document()
doc.add_heading(subject, level=1)

for group, items in grouped.items():
    doc.add_heading(group, level=2)

    for item_name, owners, update in items:
        # Bullet with item name
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item_name)
        run.bold = True
        para.paragraph_format.space_after = 0

        # Owners line
        owners_para = doc.add_paragraph(f"Owners: {owners}")
        owners_para.paragraph_format.space_after = 0

        # Weekly update line
        update_para = doc.add_paragraph(f"Weekly Update: {update}")
        update_para.paragraph_format.space_after = 6  # small space before next item

# === Save to Documents/OPP Updates ===
save_dir = os.path.join(os.path.expanduser("~"), "Documents", "OPP Updates")
os.makedirs(save_dir, exist_ok=True)

filename = f"Weekly OPP Status Update {safe_date}.docx"
filepath = os.path.join(save_dir, filename)
doc.save(filepath)

print(f"✅ Saved to: {filepath}")

import urllib.parse

# === Prepare upload ===
sharepoint_site = "cobaltsp.sharepoint.com"
site_path = "/sites/TheOneSite"
drive_path = "Company Resources Library/News & Company Updates/Weekly OPP Updates"
upload_filename = filename
upload_filepath = filepath

# === 1. Get Site ID ===
site_url = f"https://graph.microsoft.com/v1.0/sites/{sharepoint_site}:{site_path}"
site_res = requests.get(site_url, headers=headers)
site_id = site_res.json()["id"]

# === 2. Get Drive ID ===
drive_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
drive_res = requests.get(drive_url, headers=headers)
drive_id = next(
    drive["id"]
    for drive in drive_res.json()["value"]
    if drive["name"] == "Cobalt Service Partners (CSP)"
)

# === 3. Upload File ===
upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{drive_path}/{upload_filename}:/content"
with open(upload_filepath, "rb") as f:
    upload_res = requests.put(upload_url, headers=headers, data=f)

if upload_res.status_code in (200, 201):
    print("✅ File uploaded to SharePoint successfully!")
else:
    print(f"❌ Upload failed: {upload_res.status_code} - {upload_res.text}")



